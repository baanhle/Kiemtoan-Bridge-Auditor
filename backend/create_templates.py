import os
try:
    from docx import Document
except ImportError:
    import subprocess
    import sys
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "docxtpl"])
    from docx import Document

def create_tcvn_template(output_path):
    """Tạo mẫu báo cáo Word cho TCVN 11823:2017"""
    doc = Document()
    doc.add_heading('BÁO CÁO KIỂM TOÁN MẶT CẮT NGANG CẦU', 0)
    doc.add_paragraph('Tiêu chuẩn áp dụng: {{ standard_used }}')
    
    doc.add_heading('1. THÔNG SỐ CƠ BẢN ĐẦU VÀO', 1)
    
    doc.add_heading('1.1 Kích thước hình học', 2)
    doc.add_paragraph('Bề rộng nắp b = {{ input.geometry.b }} mm; Chiều cao dầm h = {{ input.geometry.h }} mm')
    doc.add_paragraph('Bề rộng sườn bw = {{ input.geometry.b_w }} mm; Chiều cao cánh hf = {{ input.geometry.h_f }} mm')
    
    doc.add_heading('1.2 Đặc trưng vật liệu', 2)
    doc.add_paragraph('Bê tông: f\'c = {{ input.materials.f_c }} MPa; Ec = {{ input.materials.E_c }} MPa')
    doc.add_paragraph('Cốt thép: fy = {{ input.materials.f_y }} MPa; Es = {{ input.materials.E_s }} MPa')
    doc.add_paragraph('Cáp DUL: fpu = {{ input.materials.f_pu }} MPa; fpy = {{ input.materials.f_py }} MPa')
    
    doc.add_heading('1.3 Bố trí cốt thép', 2)
    doc.add_paragraph('Cốt thép thường: As = {{ input.flexural_rebar.A_s }} mm2; ds = {{ input.flexural_rebar.d_s }} mm')
    doc.add_paragraph('Cáp DUL: Aps = {{ input.flexural_rebar.A_ps }} mm2; dp = {{ input.flexural_rebar.d_p }} mm')
    doc.add_paragraph('Cốt đai: Av = {{ input.shear_rebar.A_v }} mm2; s = {{ input.shear_rebar.s }} mm; Góc \u03B1 = {{ input.shear_rebar.alpha }}\u00B0')
    
    doc.add_heading('1.4 Nội lực thiết kế', 2)
    doc.add_paragraph('Moment uốn yêu cầu Mu = {{ input.forces.M_u }} kN.m')
    doc.add_paragraph('Lực cắt yêu cầu Vu = {{ input.forces.V_u }} kN')

    doc.add_heading('2. KẾT LUẬN CHUNG', 1)
    p = doc.add_paragraph('Trạng thái tổng thể mặt cắt: ')
    run = p.add_run('{{ overall_status }}')
    run.bold = True
    
    doc.add_heading('3. KIỂM TOÁN SỨC KHÁNG UỐN', 1)
    doc.add_paragraph('Tiết diện được kiểm toán dựa trên giả thiết khối ứng suất chữ nhật tương đương (Theo TCVN 11823-5:2017, Điều 5.7.2.2).')
    
    p_eq = doc.add_paragraph()
    p_eq.add_run('Công thức tính toán (Điều 5.7.3.2.2): ').italic = True
    p_eq.add_run('Mn = A_ps * f_ps * (d_p - a/2) + A_s * f_y * (d_s - a/2) - A\'_s * f\'_y * (d\'_s - a/2)')
    
    doc.add_paragraph('Chi tiết tính toán: {{ flexural.details }}')
    doc.add_paragraph('Sức kháng thiết kế (Capacity - \u03C6Mn): {{ "%.2f"|format(flexural.capacity) }} kN.m')
    doc.add_paragraph('Nội lực yêu cầu (Demand - Mu): {{ "%.2f"|format(flexural.demand) }} kN.m')
    doc.add_paragraph('Tỷ số Demand/Capacity: {{ "%.2f"|format(flexural.ratio) }} => Trạng thái: {{ "ĐẠT" if flexural.is_passed else "KHÔNG ĐẠT" }}')

    doc.add_heading('4. KIỂM TOÁN SỨC KHÁNG CẮT (MCFT)', 1)
    doc.add_paragraph('Sức kháng cắt được tính toán theo Lý thuyết trường nén biến đổi - MCFT (Theo TCVN 11823-5:2017, Điều 5.8.3.3 và 5.8.3.4.1).')
    
    p_shear = doc.add_paragraph()
    p_shear.add_run('Công thức (Điều 5.8.3.3): ').italic = True
    p_shear.add_run('Vn = Vc + Vs + Vp. Trong đó: ')
    doc.add_paragraph('Vc = 0.083 * \u03B2 * \u221Af\'c * bv * dv')
    doc.add_paragraph('Vs = [Av * fy * dv * (cot\u03B8 + cot\u03B1) * sin\u03B1] / s')
    
    doc.add_paragraph('Các thông số cắt cơ bản: {{ shear.details }}')
    doc.add_paragraph('Sức kháng thiết kế (Capacity - \u03C6Vn): {{ "%.2f"|format(shear.capacity) }} kN')
    doc.add_paragraph('Lực cắt yêu cầu (Demand - Vu): {{ "%.2f"|format(shear.demand) }} kN')
    doc.add_paragraph('Tỷ số Demand/Capacity: {{ "%.2f"|format(shear.ratio) }} => Trạng thái: {{ "ĐẠT" if shear.is_passed else "KHÔNG ĐẠT" }}')

    doc.add_heading('5. CÁC THÔNG SỐ TRUNG GIAN', 1)
    doc.add_paragraph('Chiều sâu trục trung hòa c = {{ "%.2f"|format(details.c) }} mm')
    doc.add_paragraph('Chiều sâu khối ứng suất a = {{ "%.2f"|format(details.a) }} mm')
    doc.add_paragraph('Ứng suất cáp khi phá hoại fps = {{ "%.2f"|format(details.f_ps) }} MPa')
    doc.add_paragraph('Góc nứt cắt \u03B8 = {{ "%.2f"|format(details.theta) }} độ')
    
    doc.add_page_break()
    doc.add_paragraph('Báo cáo được xuất tự động bởi Bridge Section Auditor Pro.')
    
    doc.save(output_path)
    print(f"Created template: {output_path}")

def create_eurocode_template(output_path):
    """Tạo mẫu báo cáo Word cho Eurocode 2"""
    doc = Document()
    doc.add_heading('BÁO CÁO KIỂM TOÁN MẶT CẮT CẦU', 0)
    doc.add_paragraph('Tiêu chuẩn áp dụng: {{ standard_used }} (EN 1992-1-1)')
    
    doc.add_heading('1. THÔNG SỐ CƠ BẢN ĐẦU VÀO', 1)
    
    doc.add_heading('1.1 Kích thước hình học', 2)
    doc.add_paragraph('Bề rộng nắp b = {{ input.geometry.b }} mm; Chiều cao dầm h = {{ input.geometry.h }} mm')
    doc.add_paragraph('Bề rộng sườn bw = {{ input.geometry.b_w }} mm; Chiều cao cánh hf = {{ input.geometry.h_f }} mm')
    
    doc.add_heading('1.2 Đặc trưng vật liệu', 2)
    doc.add_paragraph('Bê tông: fck = {{ input.materials.f_c }} MPa')
    doc.add_paragraph('Cốt thép: fyk = {{ input.materials.f_y }} MPa')
    doc.add_paragraph('Cáp DUL: fpuk = {{ input.materials.f_pu }} MPa')
    
    doc.add_heading('1.3 Bố trí cốt thép', 2)
    doc.add_paragraph('Cốt thép thường: As = {{ input.flexural_rebar.A_s }} mm2; ds = {{ input.flexural_rebar.d_s }} mm')
    doc.add_paragraph('Cáp DUL: Aps = {{ input.flexural_rebar.A_ps }} mm2; dp = {{ input.flexural_rebar.d_p }} mm')
    doc.add_paragraph('Cốt đai: Asw = {{ input.shear_rebar.A_v }} mm2; s = {{ input.shear_rebar.s }} mm; Góc \u03B1 = {{ input.shear_rebar.alpha }}\u00B0')
    
    doc.add_heading('1.4 Nội lực thiết kế', 2)
    doc.add_paragraph('Moment uốn MEd = {{ input.forces.M_u }} kN.m')
    doc.add_paragraph('Lực cắt VEd = {{ input.forces.V_u }} kN')

    doc.add_heading('2. KẾT LUẬN CHUNG', 1)
    doc.add_paragraph('Trạng thái tổng thể mặt cắt: {{ overall_status }}')
    
    doc.add_heading('3. KIỂM TOÁN SỨC KHÁNG UỐN (BENDING MOMENT)', 1)
    doc.add_paragraph('Phương pháp khối ứng suất tương đương (Theo EN 1992-1-1, Khoản 3.1.7).')
    
    p_ecr1 = doc.add_paragraph()
    p_ecr1.add_run('Công thức: ').italic = True
    p_ecr1.add_run('MRd = A_ps * f_pd * (d_p - \u03BBx/2) + A_s * f_yd * (d_s - \u03BBx/2)')
    
    doc.add_paragraph('Sức kháng thiết kế (MRd): {{ "%.2f"|format(flexural.capacity) }} kN.m')
    doc.add_paragraph('Nội lực yêu cầu (MEd): {{ "%.2f"|format(flexural.demand) }} kN.m')
    doc.add_paragraph('=> Trạng thái: {{ "ĐẠT" if flexural.is_passed else "KHÔNG ĐẠT" }}')

    doc.add_heading('4. KIỂM TOÁN SỨC KHÁNG CẮT (SHEAR)', 1)
    doc.add_paragraph('Sức kháng cắt theo mô hình thanh chống nghiêng có cốt đai (Variable Strut Inclination Method - Theo EN 1992-1-1, Khoản 6.2.3).')
    p_ecr2 = doc.add_paragraph()
    p_ecr2.add_run('Công thức: ').italic = True
    p_ecr2.add_run('VRd,s = (Asw / s) * z * fywd * (cot\u03B8 + cot\u03B1) * sin\u03B1')
    
    doc.add_paragraph('Chi tiết lực cắt: {{ shear.details }}')
    doc.add_paragraph('Sức kháng thiết kế (VRd): {{ "%.2f"|format(shear.capacity) }} kN')
    doc.add_paragraph('Lực cắt yêu cầu (VEd): {{ "%.2f"|format(shear.demand) }} kN')
    doc.add_paragraph('=> Trạng thái: {{ "ĐẠT" if shear.is_passed else "KHÔNG ĐẠT" }}')
    
    doc.save(output_path)
    print(f"Created template: {output_path}")

if __name__ == "__main__":
    os.makedirs('templates', exist_ok=True)
    create_tcvn_template('templates/tcvn_template.docx')
    create_eurocode_template('templates/eurocode_template.docx')
    # Copy TCVN to AASHTO
    import shutil
    shutil.copy('templates/tcvn_template.docx', 'templates/aashto_template.docx')
    print("Created template: templates/aashto_template.docx")
